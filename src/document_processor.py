import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib

# Document processing libraries
import pypdf
import pdfplumber
from docx import Document
import pandas as pd

# Text processing
import re
from datetime import datetime

class DocumentProcessor:
    """Like a universal translator documents.
    Takes any document type and converts it to clean, structured text.
    """

    def __init__(self, config: Dict[str, Any] = None):
        """Initialize the document processor"""
        self.config = config or {}
        self.supported_formats = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.csv']

        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

        # Statistics tracking
        self.stats = {
            'files_processed': 0,
            'total_pages': 0,
            'errors': 0,
            'processing_time': 0
        }

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process a single file and extract its content."""
        start_time = datetime.now()

        try:
            # Get file information
            file_info = self._get_file_info(file_path)

            # Extract content based on file type
            content = self._extract_content(file_path, file_info['extension'])

            # Clean and normalize the content
            clean_content = self._clean_content(content)

            # Extract metadata
            metadata = self._extract_metadata(file_path, content)

            processing_time = (datetime.now() - start_time).total_seconds()

            result = {
                'file_path': file_path,
                'file_info': file_info,
                'content': clean_content,
                'metadata': metadata,
                'processing_time': processing_time,
                'status': 'success',
                'success': True,  # Add this for RAG system compatibility
                'error': None
            }

            self.stats['files_processed'] += 1
            self.stats['processing_time'] += processing_time

            self.logger.info(f"✅ Processed: {Path(file_path).name}")
            return result
        
        except Exception as e:
            self.stats['errors'] += 1
            self.logger.error(f"❌ Error processing {file_path}: {str(e)}")

            return {
                'file_path': file_path,
                'content': None,
                'status': 'error',
                'success': False,  # Add this for RAG system compatibility
                'error': str(e),
                'processing_time': (datetime.now() - start_time).total_seconds()
            }
        
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information."""
        path_obj = Path(file_path)
        stat = path_obj.stat()

        return {
            'filename': path_obj.name,
            'extension': path_obj.suffix.lower(),
            'size_bytes': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'created_date': datetime.fromtimestamp(stat.st_ctime),
            'modified_date': datetime.fromtimestamp(stat.st_mtime),
            'file_hash': self._calculate_file_hash(file_path)
        }
        
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate file hash for duplicate detection."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _extract_content(self, file_path: str, extension: str) -> str:
        """Extract text content based on file type."""
        if extension == '.pdf':
            return self._extract_pdf_content(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_word_content(file_path)
        elif extension in ['.xlsx', '.xls']:
            return self._extract_excel_content(file_path)
        elif extension == '.txt':
            return self._extract_text_content(file_path)
        elif extension == '.csv':
            return self._extract_csv_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
        
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF files."""
        content = ""

        try: 
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    self.stats['total_pages'] += 1

        except Exception as e:
            self.logger.warning(f"pdfplumber failed for {file_path}, trying pypdf: {e}")

            # Fallback to pypdf
            try:
                pdf_reader = pypdf.PdfReader(file_path)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    self.stats['total_pages'] += 1

            except Exception as e2:
                raise Exception(f"Both PDF extraction methods failed: {e2}")
            
        return content
    
    def _extract_word_content(self, file_path: str) -> str:
        """Extract text from Word documents."""
        try:
            doc = Document(file_path)
            content = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    content += row_text + "\n"

            return content

        except Exception as e:
            raise Exception(f"Failed to extract Word content: {e}")
        
    def _extract_excel_content(self, file_path: str) -> str:
        """Extract content from Excel files."""
        try:
            excel_file = pd.ExcelFile(file_path)
            content = "## Excel Workbook Analysis\n\n"

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                content += f"### Sheet: {sheet_name}\n"
                content += f"This sheet contains {len(df)} rows and {len(df.columns)} columns.\n\n"
                
                # Add column definitions for better semantic understanding
                content += "#### Column Structure:\n"
                for col in df.columns:
                    # Get sample non-null values to understand data type
                    sample_vals = df[col].dropna().head(3).tolist()
                    if sample_vals:
                        sample_str = ", ".join([str(v)[:40] for v in sample_vals])
                        content += f"- {col}: {sample_str}...\n"
                    else:
                        content += f"- {col}: (empty column)\n"
                content += "\n"

                # Add data in semantic sections - create individual chunks
                content += "#### Sheet Data:\n"
                max_rows = self.config.get('excel_max_rows', 100)
                
                # Create individual semantic chunks for better retrieval
                data_chunks = []
                for idx in range(min(len(df), max_rows)):
                    row = df.iloc[idx]
                    
                    # Create natural language descriptions based on column patterns
                    natural_description = self._create_natural_row_description(row, df.columns)
                    
                    # Create a focused chunk with context
                    chunk_content = f"### Sheet: {sheet_name} - Entry {idx + 1}\n\n{natural_description}\n\n"
                    
                    # Add key fields for better searchability
                    key_fields = []
                    for col_name, value in zip(df.columns, row.values):
                        if pd.notna(value) and str(value).strip():
                            key_fields.append(f"{col_name}: {value}")
                    
                    if key_fields:
                        chunk_content += "Key fields: " + " | ".join(key_fields[:3]) + "\n"
                    
                    data_chunks.append(chunk_content)

                # Store individual chunks for better retrieval
                if hasattr(self, '_individual_chunks'):
                    self._individual_chunks.extend(data_chunks)
                else:
                    self._individual_chunks = data_chunks
                
                # Add structured content
                content += "\n".join(data_chunks)

                if len(df) > max_rows:
                    content += f"\n... ({len(df) - max_rows} additional rows not shown)\n"
                    
                content += "\n---\n\n"  # Separator between sheets

            return content
    
        except Exception as e:
            raise Exception(f"failed to extract Excel content: {e}")
    
    def _create_natural_row_description(self, row, columns):
        """Convert a data row into natural language description."""
        # Detect common data patterns and create appropriate descriptions
        
        # Personnel/Employee data pattern
        if any(col.lower() in ['name', 'employee', 'person'] for col in columns):
            return self._describe_person_row(row, columns)
        
        # Company/Organization data pattern
        elif any(col.lower() in ['company', 'organization', 'business'] for col in columns):
            return self._describe_company_row(row, columns)
        
        # Financial/Revenue data pattern
        elif any(col.lower() in ['revenue', 'sales', 'income', 'profit'] for col in columns):
            return self._describe_financial_row(row, columns)
        
        # Product/Item data pattern
        elif any(col.lower() in ['product', 'item', 'service'] for col in columns):
            return self._describe_product_row(row, columns)
        
        # Generic fallback - create readable sentence
        else:
            return self._describe_generic_row(row, columns)
    
    def _describe_person_row(self, row, columns):
        """Create natural language description for person/employee data."""
        # Find key fields
        name = self._get_field_value(row, columns, ['name', 'employee', 'person'])
        role = self._get_field_value(row, columns, ['role', 'position', 'title', 'job'])
        department = self._get_field_value(row, columns, ['department', 'dept', 'division', 'team'])
        experience = self._get_field_value(row, columns, ['experience', 'years', 'tenure'])
        
        # Build natural description
        if name:
            desc = f"{name}"
            if role:
                desc += f" works as a {role}"
            if department:
                desc += f" in the {department} department"
            if experience:
                desc += f" with {experience} of experience"
            desc += "."
            return desc
        
        return self._describe_generic_row(row, columns)
    
    def _describe_company_row(self, row, columns):
        """Create natural language description for company data."""
        company = self._get_field_value(row, columns, ['company', 'organization', 'business', 'name'])
        industry = self._get_field_value(row, columns, ['industry', 'sector', 'field'])
        revenue = self._get_field_value(row, columns, ['revenue', 'sales', 'income'])
        employees = self._get_field_value(row, columns, ['employees', 'staff', 'workforce'])
        
        if company:
            desc = f"{company}"
            if industry:
                desc += f" operates in the {industry} industry"
            if revenue:
                desc += f" with revenue of {self._format_number(revenue)}"
            if employees:
                desc += f" and employs {self._format_number(employees)} people"
            desc += "."
            return desc
        
        return self._describe_generic_row(row, columns)
    
    def _describe_financial_row(self, row, columns):
        """Create natural language description for financial data."""
        entity = self._get_field_value(row, columns, ['company', 'name', 'entity'])
        revenue = self._get_field_value(row, columns, ['revenue', 'sales', 'income'])
        profit = self._get_field_value(row, columns, ['profit', 'earnings', 'net_income'])
        
        if entity and revenue:
            desc = f"{entity} generated {self._format_number(revenue)} in revenue"
            if profit:
                desc += f" with {self._format_number(profit)} profit"
            desc += "."
            return desc
        
        return self._describe_generic_row(row, columns)
    
    def _describe_product_row(self, row, columns):
        """Create natural language description for product data."""
        product = self._get_field_value(row, columns, ['product', 'item', 'service', 'name'])
        category = self._get_field_value(row, columns, ['category', 'type', 'class'])
        price = self._get_field_value(row, columns, ['price', 'cost', 'value'])
        
        if product:
            desc = f"{product}"
            if category:
                desc += f" is a {category}"
            if price:
                desc += f" priced at {self._format_number(price)}"
            desc += "."
            return desc
        
        return self._describe_generic_row(row, columns)
    
    def _describe_generic_row(self, row, columns):
        """Create generic natural language description."""
        # Take the first few non-null values and create a sentence
        parts = []
        for col, val in zip(columns, row.values):
            if pd.notna(val) and str(val).strip():
                parts.append(f"{col} is {val}")
                if len(parts) >= 3:  # Limit to avoid overly long sentences
                    break
        
        if parts:
            return "The data shows " + ", ".join(parts) + "."
        return "Data entry with multiple fields."
    
    def _get_field_value(self, row, columns, field_names):
        """Get value for a field by trying multiple column name variations."""
        for field_name in field_names:
            for col in columns:
                if field_name.lower() in col.lower():
                    val = row[col]
                    if pd.notna(val) and str(val).strip():
                        return str(val)
        return None
    
    def _format_number(self, value):
        """Format numbers in a readable way."""
        try:
            num = float(str(value).replace(',', ''))
            if num >= 1e9:
                return f"${num/1e9:.1f} billion"
            elif num >= 1e6:
                return f"${num/1e6:.1f} million"
            elif num >= 1e3:
                return f"${num/1e3:.1f} thousand"
            else:
                return f"${num:,.0f}"
        except:
            return str(value)
        
    def _extract_text_content(self, file_path: str) -> str:
        """Extract plaint text."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
            
    def _extract_csv_content(self, file_path: str) -> str:
        """Extract CSV content."""
        try:
            df = pd.read_csv(file_path)

            # Create semantic-friendly CSV content
            content = f"## CSV Data Summary\n"
            content += f"This CSV file contains {len(df)} rows and {len(df.columns)} columns.\n\n"
            
            # Add column definitions for better semantic understanding
            content += "### Column Definitions:\n"
            for col in df.columns:
                # Get sample non-null values to understand data type
                sample_vals = df[col].dropna().head(3).tolist()
                sample_str = ", ".join([str(v)[:50] for v in sample_vals])
                content += f"- {col}: {sample_str}...\n"
            content += "\n"
            
            # Add data in semantic sections - create separate chunks per row
            content += "### Data Content:\n"
            max_rows = self.config.get('csv_max_rows', 50)
            
            # Create individual semantic chunks for better retrieval
            data_chunks = []
            for idx in range(min(len(df), max_rows)):
                row = df.iloc[idx]
                
                # Create natural language descriptions based on column patterns
                natural_description = self._create_natural_row_description(row, df.columns)
                
                # Create a focused chunk with context
                filename = os.path.basename(file_path).replace('.csv', '')
                chunk_content = f"### CSV: {filename} - Entry {idx + 1}\n\n{natural_description}\n\n"
                
                # Add key fields for better searchability
                key_fields = []
                for col_name, value in zip(df.columns, row.values):
                    if pd.notna(value) and str(value).strip():
                        key_fields.append(f"{col_name}: {value}")
                
                if key_fields:
                    chunk_content += "Key fields: " + " | ".join(key_fields[:3]) + "\n"
                
                data_chunks.append(chunk_content)

            # Store individual chunks for better retrieval (temporary approach)
            if hasattr(self, '_individual_chunks'):
                self._individual_chunks.extend(data_chunks)
            else:
                self._individual_chunks = data_chunks
            
            # For now, still return combined content but structured better
            content += "\n".join(data_chunks)
            
            if len(df) > max_rows:
                content += f"\n... ({len(df) - max_rows} additional rows not shown)\n"

            return content
        
        except Exception as e:
            raise Exception(f"Failed to extract CSV content: {e}")
        
    def _clean_content(self, content: str) -> str:
        """Clean and normalize extracted content."""
        if not content:
            return ""

        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)

        # Remove special characters that might interfere
        content = re.sub(r'[^\w\s\-.,!?:;()\[\]{}"\'/\\@#$%^&*+=<>|~`]', '', content)   
        
        # Normalize line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')

        # Remove extremely long lines (likely formatting artifacts)
        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            if len(line) < 1000:
                cleaned_lines.append(line.strip())

        return '\n'.join(cleaned_lines).strip()
    
    def _extract_metadata(self, file_path: str, content: str) -> Dict[str, Any]:
        """Extract useful metadata from the document"""
        metadata = {
            'word_count': len(content.split()) if content else 0,
            'char_count': len(content) if content else 0,
            'line_count': len(content.split('\n')) if content else 0,
        }

        # Try to extract title (first meaningful line)
        if content:
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            if lines:
                for line in lines[:5]:
                    if len(line) > 10 and len(line) < 200:
                        metadata['title'] = line
                        break

                # Extract potential keywords
                words = content.lower().split()
                word_freq = {}
                for word in words:
                    if len(word) > 4:
                        word_freq[word] = word_freq.get(word, 0) + 1

                top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
                metadata['keywords'] = [word for word, freq in top_keywords]

        return metadata
    
        
